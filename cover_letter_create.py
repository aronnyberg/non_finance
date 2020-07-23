import pandas as pd
import docx
from docx import Document

df = pd.read_csv('YOUR JOB INFO DF')
loc = 'YOUR DOCX COVER LETTER TEMPLATE'  #This should have two variables which will be changed, {descr..organization} and {descr..organization2}
end_loc = 'WHAT FOLDER IN WHICH TO SAVE FILE'
def create_cover(OGRE, ZEBRAS):
    document = Document()
    doc = docx.Document(loc)
    company_name = df[df['Unnamed: 0']==index]['Company'].values[0]
    document.add_paragraph(doc.paragraphs[0].text.replace('{company}', company_name))
    document.add_paragraph(doc.paragraphs[2].text)
    document.add_paragraph(doc.paragraphs[4].text.replace('{descr..organization}', OGRE).replace('{descr..organization2}', ZEBRAS))
    company_name_2 = company_name.replace(' ', '_')
    document.save(end_loc+'/'+company_name+'.docx')
    
number = 0 #Change this number, going down a list
index = df[df['QuickApply?']=='Yes'].iloc[0,:]['Unnamed: 0']   #Filter to only 'QuickApply?'=Yes
df[df['Unnamed: 0']==index]['Description'].values #Print description of filtered list, selecting the number above

#Here you read the cover letter and change the 


create_cover('ENTER DESCRIPTION SENTENCE 1 HERE','ENTER DESCRIPTION SENTENCE 2 HERE')
